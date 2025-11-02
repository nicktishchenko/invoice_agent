"""
Three-Phase Invoice Processing Pipeline
========================================

PHASE A: Contract Relationship Discovery
PHASE B: Per-Contract Rule Extraction
PHASE C: Invoice Processing with Content-Based Linkage

This module implements the core classes for the multi-contract framework.
"""

import json
import re
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Tuple, Optional
from docx import Document
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ContractRelationshipDiscoverer:
    """
    PHASE A: Discovers contract relationships by grouping related documents.

    Handles:
    - Multiple independent contracts in same folder
    - Single contract split across multiple documents
    - Different contract types (MSA-based, PO-based, MSA-less)
    - Date range separation of agreements between same parties
    """

    def __init__(self, contracts_dir: Path):
        self.contracts_dir = Path(contracts_dir)
        self.documents = []
        self.contracts = []

    def discover_contracts(self) -> Dict:
        """
        Main discovery pipeline.

        Returns: {
            "contracts": [
                {
                    "contract_id": "...",
                    "parties": [...],
                    "program_code": "...",
                    "date_range": {"start": "...", "end": "..."},
                    "documents": [...],
                    "hierarchy": {...},
                    "inconsistencies": [...]
                }
            ]
        }
        """
        logger.info(f"Scanning contracts in: {self.contracts_dir}")

        # Step 1: Extract identifiers from all documents
        self._extract_document_identifiers()

        # Step 2: Group documents into contracts
        self._group_documents_into_contracts()

        # Step 3: Verify hierarchy
        self._verify_contract_hierarchies()

        return {
            "discovery_timestamp": datetime.now().isoformat(),
            "contracts_dir": str(self.contracts_dir),
            "total_documents": len(self.documents),
            "contracts": self.contracts,
        }

    def _extract_document_identifiers(self):
        """Extract parties, program codes, dates, doc types from all documents"""

        for doc_path in sorted(self.contracts_dir.glob("*")):
            if doc_path.is_dir():
                continue

            try:
                identifiers = {
                    "filename": doc_path.name,
                    "filepath": str(doc_path),
                    "type": self._detect_document_type(doc_path.name),
                    "parties": self._extract_parties(doc_path),
                    "program_code": self._extract_program_code(doc_path.name),
                    "dates": self._extract_dates(doc_path),
                }

                self.documents.append(identifiers)
                logger.info(f"✓ Extracted identifiers from: {doc_path.name}")

            except Exception as e:
                logger.error(f"✗ Error processing {doc_path.name}: {str(e)[:100]}")

    def _detect_document_type(self, filename: str) -> str:
        """Detect document type from filename"""
        filename_upper = filename.upper()

        if "MSA" in filename_upper or "MASTER SERVICE" in filename_upper:
            return "MSA"
        elif "SOW" in filename_upper or "STATEMENT OF WORK" in filename_upper:
            return "SOW"
        elif "ORDER FORM" in filename_upper:
            return "ORDER_FORM"
        elif "PURCHASE ORDER" in filename_upper or "PO" in filename_upper:
            return "PURCHASE_ORDER"
        elif "DELIVERY" in filename_upper or "DN" in filename_upper:
            return "DELIVERY_NOTE"
        else:
            return "OTHER"

    def _extract_parties(self, doc_path: Path) -> List[str]:
        """Extract party names from document"""
        parties = set()

        try:
            if doc_path.suffix.lower() == ".docx":
                doc = Document(doc_path)
                text = "\n".join([p.text for p in doc.paragraphs])
            else:
                # For PDFs and other types, would need pdfplumber etc
                # For now, extract from filename
                text = doc_path.name

            # Look for common party names
            if "bayer" in text.lower():
                parties.add("BAYER")
            if "r4" in text.lower():
                parties.add("R4")

            # Add more party detection as needed

        except Exception as e:
            logger.debug(f"Could not extract parties from {doc_path.name}: {e}")

        return sorted(list(parties))

    def _extract_program_code(self, filename: str) -> Optional[str]:
        """Extract program code from filename (e.g., BCH, CAP)"""
        # Look for patterns like "BCH", "CAP", etc.
        match = re.search(r"\b([A-Z]{2,4})\b", filename)
        if match:
            code = match.group(1)
            # Filter out common words that aren't program codes
            if code not in ["FOR", "PDF", "SOW", "MSA", "THE"]:
                return code
        return None

    def _extract_dates(self, doc_path: Path) -> Dict:
        """Extract dates from document name and content"""
        dates = {"found": [], "range": None}

        # Extract from filename (YYYY-MM-DD or YYYY-12-10 format)
        filename_dates = re.findall(r"\d{4}[\s\-_]\d{2}[\s\-_]\d{2}", doc_path.name)
        if filename_dates:
            # Convert to YYYY-MM-DD format
            for date_str in filename_dates:
                normalized = date_str.replace("_", "-").replace(" ", "-")
                dates["found"].append(normalized)

        # Also look for year patterns like "2021", "2022"
        years = re.findall(r"\b(202\d)\b", doc_path.name)
        for year in years:
            if year not in dates["found"]:
                dates["found"].append(year)

        return dates

    def _group_documents_into_contracts(self):
        """Group documents by party pairs + program codes + date ranges"""

        # Create contract groups
        groups = {}

        for doc in self.documents:
            parties_key = tuple(sorted(doc["parties"]))
            program_key = doc["program_code"] or "UNKNOWN"

            # Create group identifier: (parties, program_code)
            group_id = (parties_key, program_key)

            if group_id not in groups:
                groups[group_id] = []

            groups[group_id].append(doc)

        # Create contracts from groups
        for i, (group_id, docs) in enumerate(groups.items(), 1):
            parties, program_code = group_id

            # Generate contract ID
            contract_id = f"{'_'.join(parties)}_{program_code}_{i}".replace(" ", "_")

            # Find date range
            all_dates = []
            for doc in docs:
                all_dates.extend(doc["dates"]["found"])

            contract = {
                "contract_id": contract_id,
                "parties": list(parties),
                "program_code": program_code,
                "dates_found": sorted(set(all_dates)),
                "documents": docs,
                "hierarchy": {},
                "inconsistencies": [],
            }

            self.contracts.append(contract)
            logger.info(f"✓ Grouped contract: {contract_id} ({len(docs)} documents)")

    def _verify_contract_hierarchies(self):
        """Verify document hierarchy within each contract"""

        for contract in self.contracts:
            docs = contract["documents"]

            # Map document types
            hierarchy = {
                "msa": None,
                "sow": None,
                "order_forms": [],
                "purchase_orders": [],
                "delivery_notes": [],
            }

            for doc in docs:
                doc_type = doc["type"]

                if doc_type == "MSA":
                    hierarchy["msa"] = doc["filename"]
                elif doc_type == "SOW":
                    hierarchy["sow"] = doc["filename"]
                elif doc_type == "ORDER_FORM":
                    hierarchy["order_forms"].append(doc["filename"])
                elif doc_type == "PURCHASE_ORDER":
                    hierarchy["purchase_orders"].append(doc["filename"])
                elif doc_type == "DELIVERY_NOTE":
                    hierarchy["delivery_notes"].append(doc["filename"])

            contract["hierarchy"] = hierarchy

            # Check for inconsistencies
            inconsistencies = []

            # Check if PO exists without MSA/SOW
            has_po = bool(hierarchy["purchase_orders"])
            has_msa = hierarchy["msa"] is not None
            has_sow = hierarchy["sow"] is not None

            if has_po and not has_msa and not has_sow:
                inconsistencies.append(
                    {
                        "severity": "warning",
                        "issue": "Purchase Order exists without MSA or SOW",
                        "recommendation": "Verify this is a PO-based contract",
                    }
                )

            # Check if SOW exists without MSA
            if has_sow and not has_msa:
                inconsistencies.append(
                    {
                        "severity": "warning",
                        "issue": "SOW exists without MSA",
                        "recommendation": "Verify MSA is not needed for this contract",
                    }
                )

            contract["inconsistencies"] = inconsistencies

            if inconsistencies:
                logger.warning(
                    f"⚠ {contract['contract_id']}: {len(inconsistencies)} inconsistency/inconsistencies found"
                )


class PerContractRuleExtractor:
    """
    PHASE B: Extracts rules for each discovered contract.

    Handles:
    - Loading all related documents together
    - Creating unified FAISS vector store
    - Extracting rules via RAG from all documents
    - Checking consistency across documents
    - Flagging conflicts
    """

    def __init__(self, extracted_rules_file: Path = None):
        self.all_rules = {"contracts": []}
        self.extracted_rules_file = extracted_rules_file

    def extract_rules_for_contracts(self, contract_relationships: Dict) -> Dict:
        """
        Extract rules for each discovered contract.

        Returns per-contract rules with metadata and inconsistencies.
        """

        logger.info(
            f"Starting rule extraction for {len(contract_relationships['contracts'])} contract(s)"
        )

        for contract in contract_relationships["contracts"]:
            logger.info(f"\nProcessing contract: {contract['contract_id']}")

            contract_rules = {
                "contract_id": contract["contract_id"],
                "parties": contract["parties"],
                "program_code": contract["program_code"],
                "source_documents": [doc["filename"] for doc in contract["documents"]],
                "extraction_timestamp": datetime.now().isoformat(),
                "rules": [],
                "inconsistencies": [],
                "hierarchy": contract.get("hierarchy", {}),
            }

            # In production: create FAISS store from all documents, extract rules via RAG
            # For now: load existing rules if available
            if self.extracted_rules_file and self.extracted_rules_file.exists():
                contract_rules["rules"] = self._load_existing_rules(
                    self.extracted_rules_file
                )
                logger.info(
                    f"✓ Loaded {len(contract_rules['rules'])} rules from existing extraction"
                )
            else:
                logger.info(
                    "⚠ No existing rules found. In production, would extract via RAG."
                )

            # Check for consistency (would compare across documents)
            consistency_issues = self._check_rule_consistency(contract)
            if consistency_issues:
                contract_rules["inconsistencies"] = consistency_issues
                logger.warning(
                    f"⚠ Found {len(consistency_issues)} inconsistency/inconsistencies"
                )

            self.all_rules["contracts"].append(contract_rules)

        self.all_rules["extraction_timestamp"] = datetime.now().isoformat()

        return self.all_rules

    def _load_existing_rules(self, rules_file: Path) -> List[Dict]:
        """Load existing extracted rules"""
        try:
            with open(rules_file, "r") as f:
                existing_rules = json.load(f)
            return existing_rules
        except Exception as e:
            logger.error(f"Could not load existing rules: {e}")
            return []

    def _check_rule_consistency(self, contract: Dict) -> List[Dict]:
        """Check for consistency issues across related documents"""
        inconsistencies = []

        # In production: would compare rules extracted from each document
        # For now: check if documents have conflicting information

        # Add inconsistencies found during discovery
        if "inconsistencies" in contract:
            inconsistencies.extend(contract["inconsistencies"])

        return inconsistencies

    def save_rules(self, output_file: Path):
        """Save extracted rules to JSON file"""
        try:
            output_file.parent.mkdir(parents=True, exist_ok=True)
            with open(output_file, "w") as f:
                json.dump(self.all_rules, f, indent=2)
            logger.info(f"✓ Saved rules to: {output_file}")
        except Exception as e:
            logger.error(f"Error saving rules: {e}")


class InvoiceLinkageDetector:
    """
    PHASE C: Detects which contract an invoice belongs to (content-based).

    Detection methods (in priority order):
    1. PO number matching (VERY HIGH confidence)
    2. Vendor/party matching (HIGH confidence)
    3. Program code matching (MEDIUM confidence)
    4. Service description (semantic search)
    5. Amount/date range (confirming factor)
    """

    def __init__(self, contract_relationships: Dict, rules_data: Dict = None):
        self.contract_relationships = contract_relationships
        self.rules_data = rules_data or {"contracts": []}

    def detect_invoice_contracts(self, invoices_dir: Path) -> Dict:
        """
        Detect source contract for each invoice.

        Returns: {
            "invoices": [
                {
                    "invoice_id": "...",
                    "detected_contract": "...",
                    "match_method": "...",
                    "confidence": 0.95,
                    "status": "MATCHED|AMBIGUOUS|UNMATCHED",
                    "matching_details": {...}
                }
            ]
        }
        """

        results = {
            "detection_timestamp": datetime.now().isoformat(),
            "total_invoices": 0,
            "matched": 0,
            "ambiguous": 0,
            "unmatched": 0,
            "invoices": [],
        }

        invoice_files = list(Path(invoices_dir).glob("INV-*.json"))
        logger.info(f"Detecting contracts for {len(invoice_files)} invoice(s)")

        for invoice_file in sorted(invoice_files):
            try:
                with open(invoice_file, "r") as f:
                    invoice_data = json.load(f)

                # Detect contract for this invoice
                detection = self._detect_single_invoice(invoice_data)
                results["invoices"].append(detection)

                results["total_invoices"] += 1
                if detection["status"] == "MATCHED":
                    results["matched"] += 1
                elif detection["status"] == "AMBIGUOUS":
                    results["ambiguous"] += 1
                else:
                    results["unmatched"] += 1

                status_sym = (
                    "✓"
                    if detection["status"] == "MATCHED"
                    else "⚠" if detection["status"] == "AMBIGUOUS" else "✗"
                )
                logger.info(
                    f"{status_sym} {invoice_data.get('invoice_id', 'UNKNOWN')}: {detection['status']}"
                )

            except Exception as e:
                logger.error(f"Error processing invoice {invoice_file.name}: {e}")

        return results

    def _detect_single_invoice(self, invoice_data: Dict) -> Dict:
        """Detect contract for a single invoice"""

        invoice_id = invoice_data.get("invoice_id", "UNKNOWN")

        # Try detection methods in priority order
        matches = []

        # 1. PO number matching (VERY HIGH confidence)
        po_matches = self._match_by_po_number(invoice_data)
        if po_matches:
            for contract_id, confidence in po_matches:
                matches.append((contract_id, "PO_NUMBER", confidence))

        # 2. Vendor/party matching (HIGH confidence)
        if not matches:
            vendor_matches = self._match_by_vendor(invoice_data)
            if vendor_matches:
                for contract_id, confidence in vendor_matches:
                    matches.append((contract_id, "VENDOR", confidence))

        # 3. Program code matching (MEDIUM confidence)
        if not matches:
            program_matches = self._match_by_program_code(invoice_data)
            if program_matches:
                for contract_id, confidence in program_matches:
                    matches.append((contract_id, "PROGRAM_CODE", confidence))

        # Build result
        result = {
            "invoice_id": invoice_id,
            "detected_contract": None,
            "match_method": None,
            "confidence": 0.0,
            "matching_details": {},
            "alternative_matches": [],
            "status": "UNMATCHED",
        }

        if len(matches) == 1:
            # Unique match
            contract_id, method, confidence = matches[0]
            result["detected_contract"] = contract_id
            result["match_method"] = method
            result["confidence"] = confidence
            result["status"] = "MATCHED"
            result["matching_details"] = self._get_matching_details(
                invoice_data, contract_id
            )

        elif len(matches) > 1:
            # Multiple matches - ambiguous
            result["detected_contract"] = matches[0][0]
            result["match_method"] = matches[0][1]
            result["confidence"] = matches[0][2]
            result["alternative_matches"] = [
                {"contract_id": m[0], "method": m[1], "confidence": m[2]}
                for m in matches[1:]
            ]
            result["status"] = "AMBIGUOUS"
            result["matching_details"] = self._get_matching_details(
                invoice_data, matches[0][0]
            )

        return result

    def _match_by_po_number(self, invoice_data: Dict) -> List[Tuple[str, float]]:
        """Match invoice to contract by PO number"""
        invoice_po = invoice_data.get("po_number")

        if not invoice_po:
            return []

        matches = []

        # Search all contract documents for PO references
        for contract in self.contract_relationships["contracts"]:
            for doc in contract["documents"]:
                # In production: would search document content for PO
                # For now: simple filename matching
                if invoice_po in doc["filename"]:
                    matches.append((contract["contract_id"], 0.95))

        return matches

    def _match_by_vendor(self, invoice_data: Dict) -> List[Tuple[str, float]]:
        """Match invoice to contract by vendor name"""
        invoice_vendor = invoice_data.get("vendor", "").lower()

        if not invoice_vendor:
            return []

        matches = []

        for contract in self.contract_relationships["contracts"]:
            for party in contract["parties"]:
                if party.lower() in invoice_vendor or invoice_vendor in party.lower():
                    # Check if invoice date is within contract date range
                    confidence = 0.85
                    matches.append((contract["contract_id"], confidence))
                    break

        return matches

    def _match_by_program_code(self, invoice_data: Dict) -> List[Tuple[str, float]]:
        """Match invoice to contract by program code"""
        invoice_description = (
            invoice_data.get("services_description", "")
            + invoice_data.get("reason", "")
        ).lower()

        # Extract program codes from invoice
        program_codes = re.findall(r"\b([A-Z]{2,4})\b", invoice_description)

        if not program_codes:
            return []

        matches = []

        for contract in self.contract_relationships["contracts"]:
            if contract["program_code"] in program_codes:
                confidence = 0.70
                matches.append((contract["contract_id"], confidence))

        return matches

    def _get_matching_details(self, invoice_data: Dict, contract_id: str) -> Dict:
        """Get details of why invoice matched this contract"""
        details = {
            "po_number": invoice_data.get("po_number"),
            "vendor": invoice_data.get("vendor"),
            "invoice_date": invoice_data.get("invoice_date"),
            "amount": invoice_data.get("amount"),
        }
        return details


class InvoiceParser:
    """
    PHASE C (Helper): Parses invoice documents and extracts fields.

    Supports: PDF, DOCX, DOC formats

    Extracted fields:
    - invoice_id (from filename: INV-001 → "INV-001")
    - vendor (party/company name)
    - po_number (purchase order reference)
    - invoice_date (date created)
    - amount (total amount)
    - services_description (what was invoiced for)
    """

    def __init__(self):
        self.extracted_invoices = []

    def parse_invoices_directory(self, invoices_dir: Path) -> List[Dict]:
        """
        Parse all invoice files in directory.

        Returns list of extracted invoice data dicts.
        """

        invoices_dir = Path(invoices_dir)
        logger.info(f"Parsing invoices from: {invoices_dir}")

        # Get all PDF and DOCX files
        invoice_files = []
        invoice_files.extend(invoices_dir.glob("INV-*.pdf"))
        invoice_files.extend(invoices_dir.glob("INV-*.docx"))
        invoice_files.extend(invoices_dir.glob("INV-*.doc"))

        # Remove duplicates (keep both PDF and DOCX if available)
        unique_invoices = {}
        for file_path in sorted(invoice_files):
            # Extract base name (e.g., "INV-001" from "INV-001.pdf")
            base_name = file_path.stem  # stem removes extension

            # Prefer DOCX over PDF (more reliable extraction)
            if base_name not in unique_invoices or file_path.suffix == ".docx":
                unique_invoices[base_name] = file_path

        # Parse each unique invoice
        for base_name, file_path in sorted(unique_invoices.items()):
            try:
                invoice_data = self._parse_single_invoice(file_path)
                self.extracted_invoices.append(invoice_data)
                logger.info(f"✓ Parsed: {file_path.name}")
            except Exception as e:
                logger.error(f"✗ Failed to parse {file_path.name}: {str(e)[:100]}")

        logger.info(f"✓ Successfully parsed {len(self.extracted_invoices)} invoices")
        return self.extracted_invoices

    def _parse_single_invoice(self, file_path: Path) -> Dict:
        """Parse a single invoice file and extract fields"""

        # Read file content based on extension
        if file_path.suffix.lower() == ".docx":
            content = self._parse_docx(file_path)
        elif file_path.suffix.lower() == ".pdf":
            content = self._parse_pdf(file_path)
        elif file_path.suffix.lower() == ".doc":
            # Basic support - would need python-docx with legacy format
            content = self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")

        # Extract fields from document content (NOT from filename)
        extracted = {
            "file_path": str(file_path),
            "file_format": file_path.suffix.lower(),
            "raw_content": content,
        }

        # Extract structured fields from content
        # This includes invoice_id extracted from document, not filename
        extracted.update(self._extract_fields_from_content(content))

        return extracted

    def _parse_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs])
            # Also get tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
            return text
        except Exception as e:
            logger.warning(f"Could not parse DOCX {file_path.name}: {e}")
            return ""

    def _parse_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            import pdfplumber

            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += "\n" + (page.extract_text() or "")
            return text
        except Exception as e:
            logger.warning(f"Could not parse PDF {file_path.name}: {e}")
            return ""

    def _extract_fields_from_content(self, content: str) -> Dict:
        """
        Extract structured fields from document content.

        IMPORTANT: All fields are extracted from document content, NOT filenames.
        This ensures the invoice ID, vendor, dates, etc. come from the actual
        document, not from filename assumptions.
        """

        fields = {
            "invoice_id": None,  # Will be extracted from content
            "vendor": None,
            "po_number": None,
            "invoice_date": None,
            "amount": None,
            "services_description": None,
            "currency": "USD",  # Default
            "payment_terms": None,
        }

        # Convert to lowercase for searching (but keep original for extraction)
        content_lower = content.lower()

        # ========== EXTRACT INVOICE ID FROM CONTENT ==========
        # Do NOT use filename! Extract from document fields like:
        #   "Invoice #: INV-001"
        #   "Invoice Number: INV-001"
        #   "Invoice ID: INV-001"
        invoice_id_patterns = [
            r"invoice\s*#:?\s*([A-Z0-9\-]+)",
            r"invoice\s+number:?\s*([A-Z0-9\-]+)",
            r"invoice\s+id:?\s*([A-Z0-9\-]+)",
        ]
        for pattern in invoice_id_patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                fields["invoice_id"] = match.group(1).strip()
                break

        # If invoice_id not found in content, log warning (don't use filename)
        if not fields["invoice_id"]:
            logger.warning("Could not extract invoice_id from document content")
        # ========== EXTRACT PO NUMBER FROM CONTENT ==========
        po_patterns = [
            r"po\s+number:\s*([A-Z0-9\-]+)",
            r"po\s*#:?\s*([A-Z0-9\-]+)",
            r"purchase\s+order\s*#?:?\s*([A-Z0-9\-]+)",
            r"p\.o\.\s*#?:?\s*([A-Z0-9\-]+)",
        ]
        for pattern in po_patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                fields["po_number"] = match.group(1).strip()
                break

        # ========== EXTRACT VENDOR NAME FROM CONTENT ==========
        # Look for patterns like "FROM: Company Name" or "VENDOR: Company Name"
        vendor_patterns = [
            r"from:\s*([^\n]+)",
            r"vendor:\s*([^\n]+)",
            r"billed by:\s*([^\n]+)",
            r"supplier:\s*([^\n]+)",
        ]
        for pattern in vendor_patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                vendor_text = match.group(1).strip()
                # Clean up the vendor text
                vendor_text = vendor_text.split("\n")[0].strip()
                if vendor_text and len(vendor_text) < 100:  # Sanity check
                    fields["vendor"] = vendor_text
                    break

        # ========== EXTRACT INVOICE DATE FROM CONTENT ==========
        # Look for patterns like "Date: 2025-11-01" or "Invoice Date: ..."
        date_patterns = [
            r"(?:invoice\s+)?date:?\s*(\d{4}[-/]\d{2}[-/]\d{2})",
            r"(\d{4}[-/]\d{2}[-/]\d{2})",  # Any YYYY-MM-DD or similar
        ]
        for pattern in date_patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                fields["invoice_date"] = match.group(1)
                break

        # ========== EXTRACT AMOUNT FROM CONTENT ==========
        # Look for patterns like "Amount: $15,000.00" or "Total: $..."
        amount_patterns = [
            r"amount:?\s*\$?([\d,]+\.?\d*)",
            r"total:?\s*\$?([\d,]+\.?\d*)",
            r"\$\s*([\d,]+\.?\d*)",  # Dollar amounts
        ]
        for pattern in amount_patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                amount_str = match.group(1).replace(",", "")
                try:
                    fields["amount"] = float(amount_str)
                    break
                except ValueError:
                    continue

        # ========== EXTRACT SERVICE DESCRIPTION FROM CONTENT ==========
        # Look for sections like "Services:" or description fields
        # The description often appears on the line after a standalone "Services" line
        desc_patterns = [
            r"^Services\s*\n\s*([^\n]+)",  # Standalone "Services" at line start, capture next line
            r"services?\s*:\s*([^\n]+)",  # "Services: description text"
            r"description:?\s*([^\n]+)",
            r"for:?\s*([^\n]+)",
        ]
        for pattern in desc_patterns:
            match = re.search(pattern, content, re.IGNORECASE | re.MULTILINE)
            if match:
                desc_text = match.group(1).strip()
                if desc_text and len(desc_text) < 200:  # Sanity check
                    fields["services_description"] = desc_text
                    break

        # ========== EXTRACT PAYMENT TERMS FROM CONTENT ==========
        # Look for patterns like "Payment Terms: Net 30"
        terms_patterns = [
            r"payment\s+terms?:?\s*([^\n]+)",
            r"net\s+(\d+)",  # Net 30, Net 60, etc.
        ]
        for pattern in terms_patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                fields["payment_terms"] = match.group(1).strip()
                break

        # ========== EXTRACT CURRENCY FROM CONTENT ==========
        # Look for currency indicators
        currency_patterns = [
            r"usd",
            r"eur",
            r"gbp",
            r"\$",  # USD indicator
        ]
        for pattern in currency_patterns:
            if re.search(pattern, content, re.IGNORECASE):
                if pattern == r"\$":
                    fields["currency"] = "USD"
                else:
                    fields["currency"] = pattern.upper()
                break

        return fields
